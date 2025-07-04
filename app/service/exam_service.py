import os
import re
from http import HTTPStatus
from typing import List
from docx import Document
import zipfile

from app.core.app_config import CustomException
from app.model.models import Subject
from app.schema.schema import ExamPreview, QuestionPreview, QuestionChoice
from app.utils.date_time_utils import is_valid_date, convert_to_datetime
from app.model.models import Exam, ExamQuestion, QuestionChoice
from app.core.db_connect import SessionLocal
import uuid

IMAGE_FOLDER = "static/images"
db = SessionLocal()


def extract_exam_info(docx_path: str) -> ExamPreview:
    db = SessionLocal()
    doc = Document(docx_path)
    exam = ExamPreview()

    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    patterns = {
        "subject_code": r"Subject:\s*(.+)",
        "number_quiz": r"Number of Quiz:\s*(\d+)",
        "description": r"Lecturer:\s*(.+)",
        "start_date": r"Date:\s*([\d\w\-\/]+)",
    }
    subject_code = None

    for field, pattern in patterns.items():
        match = re.search(pattern, full_text, flags=re.IGNORECASE)
        if not match:
            continue

        value = match.group(1).strip()

        if field == "number_quiz":
            setattr(exam, field, int(value))
        elif field == "start_date":
            if value == "YYYY-MM-DD":
                exam.start_date = None
            elif is_valid_date(value, "%d-%m-%Y"):
                exam.start_date = value
            else:
                exam.start_date = None
        elif field == "subject_code":
            subject_code = value
        else:
            setattr(exam, field, value)

    # Query DB for subject_id if subject_code is found
    if subject_code:
        subject = db.query(Subject).filter(Subject.subject_code == subject_code).first()
        exam.subject_id = str(subject.id) if subject else None
    else:
        exam.subject_id = None

    return exam


def save_image_from_docx(docx_path: str) -> dict:
    image_map = {}
    with zipfile.ZipFile(docx_path) as docx_zip:
        for file in docx_zip.namelist():
            if file.startswith("word/media/"):
                ext = os.path.splitext(file)[1]
                img_data = docx_zip.read(file)
                filename = f"{uuid.uuid4()}{ext}"
                filepath = os.path.join(IMAGE_FOLDER, filename)
                with open(filepath, "wb") as f:
                    f.write(img_data)
                image_map[file.split("/")[-1]] = filepath
    return image_map


def extract_questions(docx_path: str) -> List[QuestionPreview]:
    doc = Document(docx_path)
    questions: List[QuestionPreview] = []
    image_map = save_image_from_docx(docx_path)
    image_list = list(image_map.values())
    image_index = 0

    for table in doc.tables:
        q = QuestionPreview(options=[])
        correct_letter = None

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print("ROW CELLS:", cells)  # Debugging line

            key, value = cells[0], cells[1]

            # Bỏ qua dòng QN
            if re.match(r"^QN=\d+", key, re.IGNORECASE):
                q.content = value.strip()
                continue

            # Đáp án đúng
            if re.match(r"^ANSWER\s*:?", key, re.IGNORECASE):
                correct_letter = value.strip().upper()
                continue

            # MARK
            if re.match(r"^MARK\s*:?", key, re.IGNORECASE):
                try:
                    q.mark = float(value.strip())
                except ValueError:
                    q.mark = 0.0
                continue

            # UNIT
            if re.match(r"^UNIT\s*:?", key, re.IGNORECASE):
                q.unit = value.strip()
                continue

            # MIX CHOICES
            if re.match(r"^MIX CHOICES\s*:?", key, re.IGNORECASE):
                q.mix_choices = value.strip().lower() == "yes"
                continue

            # Đáp án lựa chọn
            if re.match(r"^[a-dA-D]\.", key):
                q.options.append(
                    QuestionChoice(content=value.strip(), is_correct=False)
                )
                continue

            # Gán nội dung câu hỏi
            if not q.content:
                q.content = value
                continue

        # Gán hình nếu có
        if "file:" in (q.content or "").lower() and image_index < len(image_list):
            q.file_id = image_list[image_index]
            image_index += 1

        # Gán đáp án đúng
        if correct_letter:
            idx = ord(correct_letter) - ord("A")
            if 0 <= idx < len(q.options):
                q.options[idx].is_correct = True

        # Chỉ thêm nếu có nội dung hoặc ảnh và đáp án
        if (q.content or q.file_id) and q.options:
            questions.append(q)

    return questions


def parse_full_exam(docx_path: str) -> ExamPreview:
    exam = extract_exam_info(docx_path)
    exam.questions = extract_questions(docx_path)
    return exam


def save_exam(exam: ExamPreview):
    db = SessionLocal()
    try:
        # Validate exam data
        if not exam.subject_id:
            raise CustomException(
                f"Môn học không được cung cấp.",
                HTTPStatus.BAD_REQUEST
            )

        subject = db.query(Subject).filter(Subject.id == int(exam.subject_id)).first()
        if not subject:
            raise CustomException(
                f"Môn học có ID = {exam.subject_id} không tồn tại trong hệ thống.",
                HTTPStatus.NOT_FOUND
            )

        if not exam.questions or len(exam.questions) == 0:
            raise CustomException("Chưa có câu hỏi nào cho bài kiểm tra này.", HTTPStatus.BAD_REQUEST)

        if len(exam.questions) < exam.number_quiz:
            raise CustomException(
                f"Số câu hỏi chưa đủ: {len(exam.questions)} / {exam.number_quiz}.",
                HTTPStatus.BAD_REQUEST
            )

        if not exam.duration_minutes or exam.duration_minutes <= 0:
            raise CustomException("Chưa thiết lập thời gian cho bài thi.", HTTPStatus.BAD_REQUEST)

        if not exam.start_date:
            raise CustomException("Ngày bắt đầu không được để trống.", HTTPStatus.BAD_REQUEST)

        # 1. Tạo exam
        new_exam = Exam(
            name=f"{subject.subject_code}_{uuid.uuid4().hex[:6].upper()}" if not exam.name else exam.name,
            subject_id=int(exam.subject_id),
            number_quiz=exam.number_quiz,
            start_date=convert_to_datetime(exam.start_date),
            duration_minutes=exam.duration_minutes,
            description=exam.description or "",
            is_active=exam.is_active or True
        )
        db.add(new_exam)
        db.commit()
        db.refresh(new_exam)  # Lấy lại ID

        # 2. Thêm câu hỏi
        for question in exam.questions:
            new_question = ExamQuestion(
                exam_id=new_exam.id,
                content=question.content or "",
                file_id=question.file_id,
                mark=question.mark,
                unit=question.unit,
                mix_choices=question.mix_choices
            )
            db.add(new_question)
            db.flush()

            # 3. Thêm đáp án
            for option in question.options:
                # Use the model QuestionChoice, not the schema
                new_choice = app.model.models.QuestionChoice(
                    question_id=new_question.id,
                    content=option.content,
                    is_correct=option.is_correct
                )
                db.add(new_choice)

        db.commit()
        return new_exam.id

    except CustomException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        raise CustomException(f"Đã xảy ra lỗi hệ thống khi lưu bài thi: {str(e)}", HTTPStatus.INTERNAL_SERVER_ERROR)
    finally:
        db.close()